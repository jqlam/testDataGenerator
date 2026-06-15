# Input
#   file prefix, if not specify randomly generated
#   file size, if not specify default 2KB
#   destination path

# Output
#   DOCXFile

from docx import Document
from docx.shared import Inches
from util.randomObjectName import randomObjectName
import logging
import os
logger = logging.getLogger(__name__)

def generateDOCX(filePrefix, fileSize, destinationPath, baseDir, name=None, debug=False):
    extension = ".docx"
    lorem_path = f"{baseDir}/resources/loremIpsum.txt"

    loremSize = 2733
    baseSize = 36563 #The smallest size a word doc can be
    size = fileSize - baseSize

    with open(lorem_path, "r") as loremfile:
        # Creates string to write into output
        # lorem is 1 KB
        lorem = loremfile.read()
        loremRemainder = loremfile.read((size * 2)%loremSize)

    # Initialize a new document object
    doc = Document()
    

    # GenerateFileName
    fileName = ""
    if filePrefix is not None:
        fileName += str(filePrefix) + "_"
    if not name:
        fileName += randomObjectName()
    else:
        fileName += name
    fileName += extension

    # Generate Path
    path = ""
    if destinationPath is not None:
        path += destinationPath
        if path[:-1] != '/':
            path += "/"
    else:
        path += f"{baseDir}/output/"
    path += fileName

    # Add a main heading (Level 0)
    doc.add_heading(fileName[:-5], level=0)

    # Add a standard text paragraph
    size = 0
    prevSize = 0
    difference = 0
    while size < fileSize:
        p = doc.add_paragraph(lorem)
        doc.save(path)
        prevSize = size
        size = os.path.getsize(path)
        difference = size - prevSize
        if size < fileSize:
            loops = int(((fileSize)-size)/difference)
            for i in range(loops):
                p = doc.add_paragraph(lorem)
            doc.save(path)
            size = os.path.getsize(path)
        if debug:
            print(f"Genererating {fileName} | Expected size: {fileSize} | Current Size: {size}")
            logger.info(f"Genererating {fileName} | Expected size: {fileSize} | Current Size: {size}")

    print(f"Generated {path}")
    logger.info(f"Generated {path}")